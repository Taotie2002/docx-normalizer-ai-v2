"""
DocxExtractor - 脱骨解析器
SRS-2026-002 V11.2 | Phase 1

将 .docx 文档彻底"脱骨"，压扁为 DocumentIRBlock 序列。
目标：纯文本流的零丢失提取，为后续分类和重构做准备。
"""

import uuid
from docx import Document
from docx.oxml.ns import qn
from typing import List, Optional
import logging

from src.core.ir_block import DocumentIRBlock, BlockLabel
from src.core.exceptions import ExtractionError

logger = logging.getLogger(__name__)


class DocxExtractor:
    """
    脱骨解析器：将 .docx 文档压扁为 IR Block 序列
    
    Phase 1 目标：
    - 提取纯文本段落
    - 保证 source_para_idx 严格递增（零丢失审计基础）
    - 初步识别标题样式
    - 暂不处理复杂对象（表格、图片等）
    """
    
    def __init__(self, file_path: str):
        """
        初始化脱骨器
        
        Args:
            file_path: .docx 文件路径
        """
        self.file_path = file_path
        self._doc: Optional[Document] = None
        self._total_paragraphs: int = 0
    
    @property
    def doc(self) -> Document:
        """懒加载 Document 对象"""
        if self._doc is None:
            self._doc = Document(self.file_path)
            self._total_paragraphs = len(self._doc.paragraphs)
        return self._doc
    
    def extract_to_ir(self) -> List[DocumentIRBlock]:
        """
        将文档彻底脱骨，压扁为 IR Block 序列
        
        Phase 1 策略：
        - 保留所有非空段落
        - label 默认为 UNPROCESSED，后续由 Classifier 填充
        - 空段落可选保留（用于保持文档结构）
        
        Returns:
            List[DocumentIRBlock]: IR Block 序列
            
        Raises:
            ExtractionError: 提取失败
        """
        ir_blocks: List[DocumentIRBlock] = []
        skipped_empty: int = 0
        
        try:
            for idx, para in enumerate(self.doc.paragraphs):
                raw_text = para.text
                
                # Phase 1: 跳过纯空段落（仅空白字符）
                if not raw_text.strip():
                    skipped_empty += 1
                    logger.debug(f"[Extractor] 跳过空段落 idx={idx}")
                    continue
                
                # 创建 IR Block
                block = self._create_block(
                    para=para,
                    raw_text=raw_text,
                    source_idx=idx
                )
                
                ir_blocks.append(block)
            
            logger.info(
                f"[Extractor] 脱骨完成: {len(ir_blocks)} 个有效段落, "
                f"{skipped_empty} 个空段落被跳过"
            )
            
            # V-03 验证：检查 source_para_idx 是否严格递增
            self._validate_sequence(ir_blocks)
            
            return ir_blocks
            
        except Exception as e:
            raise ExtractionError(f"提取失败: {e}") from e
    
    def _create_block(
        self, 
        para, 
        raw_text: str, 
        source_idx: int
    ) -> DocumentIRBlock:
        """
        创建单个 DocumentIRBlock
        
        Args:
            para: python-docx Paragraph 对象
            raw_text: 段落原始文本
            source_idx: 原始段落索引
            
        Returns:
            DocumentIRBlock
        """
        block = DocumentIRBlock(
            block_id=str(uuid.uuid4()),
            text=raw_text,
            label=BlockLabel.UNPROCESSED,  # Phase 1 默认值
            confidence=1.0,
            source_para_idx=source_idx,
            classifier_source="SYSTEM"
        )
        
        # 提取物理特征：识别标题样式
        self._extract_heading_level(block, para)
        
        return block
    
    def _extract_heading_level(self, block: DocumentIRBlock, para) -> None:
        """
        提取标题级别
        
        python-docx 的 Heading 样式命名规范：
        - "Heading 1", "Heading 2", ..., "Heading 9"
        - Word 可能使用本地化名称，如"标题 1"
        
        Args:
            block: IR Block
            para: 段落对象
        """
        style_name = para.style.name if para.style else ""
        
        # 匹配 Heading 样式
        if style_name.startswith("Heading"):
            try:
                # 尝试解析 "Heading N" 格式
                level_str = style_name.replace("Heading", "").strip()
                if level_str.isdigit():
                    block.heading_level = int(level_str)
                    block.is_list_item = False
                    logger.debug(
                        f"[Extractor] 检测到标题级别 {block.heading_level}: "
                        f"{block.text[:30]}"
                    )
            except ValueError:
                pass
        
        # 备选：检测中文"标题"样式
        elif style_name.startswith("标题"):
            try:
                level_str = style_name.replace("标题", "").strip()
                if level_str.isdigit():
                    block.heading_level = int(level_str)
            except ValueError:
                pass
    
    def _validate_sequence(self, blocks: List[DocumentIRBlock]) -> None:
        """
        V-03 验证：检查 source_para_idx 是否严格递增
        
        Args:
            blocks: IR Block 序列
            
        Raises:
            ExtractionError: 序列不连续
        """
        if not blocks:
            return
            
        for i in range(1, len(blocks)):
            if blocks[i].source_para_idx <= blocks[i-1].source_para_idx:
                raise ExtractionError(
                    f"V-03 序列验证失败: "
                    f"block[{i}] source_para_idx={blocks[i].source_para_idx} "
                    f"<= block[{i-1}] source_para_idx={blocks[i-1].source_para_idx}"
                )
        
        logger.debug(
            f"[Extractor] V-03 序列验证通过: "
            f"{blocks[0].source_para_idx} -> {blocks[-1].source_para_idx}"
        )
    
    def get_document_info(self) -> dict:
        """
        获取文档基本信息（调试用）
        
        Returns:
            dict: 文档统计信息
        """
        return {
            "file_path": self.file_path,
            "total_paragraphs": self._total_paragraphs,
            "is_loaded": self._doc is not None
        }


# === 便捷函数 ===

def extract_file(file_path: str) -> List[DocumentIRBlock]:
    """
    便捷函数：一行代码提取 .docx 文件
    
    Args:
        file_path: .docx 文件路径
        
    Returns:
        List[DocumentIRBlock]: IR Block 序列
    """
    extractor = DocxExtractor(file_path)
    return extractor.extract_to_ir()
