"""
DocxCompiler - 重构编译器
SRS-2026-002 V11.2 | Phase 1

消费 IR 序列与 DFGP 规范，在空白 Document 中灌注内容。
Phase 1: 仅保证纯文本流零丢失，不应用格式。
"""

import uuid
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Cm, Twips
from typing import List, Optional
import logging

from src.core.ir_block import DocumentIRBlock
from src.core.exceptions import CompilationError

logger = logging.getLogger(__name__)


class DocxCompiler:
    """
    重构编译器：将 IR Block 序列灌注为空白 Document
    
    Phase 1 目标：
    - 创建空白容器
    - 根据 heading_level 自动选择 add_heading 或 add_paragraph
    - 保证文本内容 100% 一致（V-03 零丢失）
    - 暂不应用格式（Phase 2 由 DFGP 接管）
    """
    
    def __init__(self):
        """创建空白容器"""
        self.doc = Document()
    
    def build_from_ir(
        self, 
        ir_blocks: List[DocumentIRBlock], 
        output_path: str,
        apply_pagination: bool = False  # Phase 2 启用
    ) -> None:
        """
        将 IR 序列重新灌注为物理文档
        
        Args:
            ir_blocks: IR Block 序列
            output_path: 输出文件路径
            apply_pagination: 是否应用分页标签 (Phase 2)
            
        Raises:
            CompilationError: 重构失败
        """
        if not ir_blocks:
            raise CompilationError("IR Block 序列为空")
        
        logger.info(f"[Compiler] 准备重构，收到 {len(ir_blocks)} 个 IR Blocks")
        
        try:
            for block in ir_blocks:
                self._add_block_to_doc(block)
            
            # Phase 1.5: 可选分页标签注入
            if apply_pagination:
                self._apply_pagination_tags(ir_blocks)
            
            self.doc.save(output_path)
            logger.info(f"[Compiler] 重构完成！已保存至: {output_path}")
            
        except Exception as e:
            raise CompilationError(f"重构失败: {e}") from e
    
    def _add_block_to_doc(self, block: DocumentIRBlock) -> None:
        """
        将单个 IR Block 灌注到 Document
        
        Args:
            block: IR Block
        """
        text = block.text
        
        if not text:
            return
        
        # Phase 1 策略：根据 heading_level 选择添加方式
        if block.heading_level is not None and block.heading_level > 0:
            # 标题段落
            level = min(block.heading_level, 9)  # Word 标题最多9级
            para = self.doc.add_heading(text, level=level)
        else:
            # 普通段落
            para = self.doc.add_paragraph(text)
    
    def _apply_pagination_tags(self, ir_blocks: List[DocumentIRBlock]) -> None:
        """
        Phase 2: 应用分页标签
        
        将 IR Block 的 pagination 字段转化为 OOXML 底层标签：
        - <w:keepNext/>: 与下段同页
        - <w:pageBreakBefore/>: 段前分页
        - <w:outlineLvl/>: 大纲级别
        
        Args:
            ir_blocks: IR Block 序列
        """
        for block, para in zip(ir_blocks, self.doc.paragraphs):
            pPr = para._element.find(
                qn('w:pPr')
            ) or self._create_pPr(para)
            
            # keepWithNext
            if block.pagination.get("keep_with_next"):
                keepNext = OxmlElement('w:keepNext')
                pPr.append(keepNext)
            
            # pageBreakBefore
            if block.pagination.get("page_break_before"):
                pageBreak = OxmlElement('w:pageBreakBefore')
                pPr.append(pageBreak)
            
            # outlineLvl (大纲级别)
            if block.heading_level is not None:
                outlineLvl = OxmlElement('w:outlineLvl')
                outlineLvl.set(qn('w:val'), str(block.heading_level - 1))
                pPr.append(outlineLvl)
    
    def _create_pPr(self, para) -> OxmlElement:
        """创建段落属性节点"""
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
        return pPr


# === 便捷函数 ===

def compile_ir_to_file(
    ir_blocks: List[DocumentIRBlock], 
    output_path: str
) -> None:
    """
    便捷函数：将 IR Block 序列直接编译为文件
    
    Args:
        ir_blocks: IR Block 序列
        output_path: 输出文件路径
    """
    compiler = DocxCompiler()
    compiler.build_from_ir(ir_blocks, output_path)
