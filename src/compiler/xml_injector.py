"""
SemanticCompiler - 语义重构编译器（纯样式方案）
SRS-2026-002 V11.2 | Phase 2

核心职责：
1. 根据 DocumentIRBlock 的 label，通过 DFGP 获取格式参数
2. 使用 python-docx 高级 API 设置样式（paragraph.style, run.font.name）
3. 对于 Word 内置样式不满足的需求，通过底层 XML 设置 outlineLevel

纯样式方案 vs XML 注入方案：
- XML 注入：直接操作 <w:rPr>/<w:pPr> 节点
- 纯样式：通过 paragraph.style + run.font.name + paragraph.alignment
- 本实现采用混合方案：优先使用 python-docx API，必要时保留底层 XML 控制

关键实现：
- _apply_paragraph_style(): 使用 paragraph.style 设置段落样式
- _apply_run_font(): 通过 run.font.name 设置字体
- _apply_paragraph_format(): 通过 paragraph.alignment/indentation 设置段落格式
- _ensure_outline_level(): 确保 outlineLevel 设置（TOC 生成必需）
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Twips
from docx.enum.style import WD_STYLE_TYPE
from typing import List, Optional
import logging

from src.core.ir_block import DocumentIRBlock
from src.core.dfgp_manager import DFGPManager, StyleParams
from src.core.exceptions import CompilationError
from src.extractor.toc_detector import TocInfo
from src.compiler.toc_generator import TocGenerator

logger = logging.getLogger(__name__)


# Word 内置样式名称映射
BUILTIN_STYLE_NAMES = {
    'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
    'Heading 5', 'Heading 6', 'Heading 7', 'Heading 8', 'Heading 9',
    'Normal', 'List Bullet', 'List Number', 'Title', 'Subtitle',
    'Quote', 'Caption', 'Footnote Text', 'Header', 'Footer'
}


class SemanticCompiler:
    """
    语义重构编译器（纯样式方案）
    
    使用 python-docx 高级 API 设置格式，仅在必要时操作底层 XML。
    """
    
    def __init__(self, dfgp_manager: Optional[DFGPManager] = None):
        """
        初始化语义编译器
        
        Args:
            dfgp_manager: 可选 DFGP 管理器，默认使用 GB/T 9704 标准
        """
        self.doc = Document()
        self.dfgp = dfgp_manager or DFGPManager()
        
        # 设置默认页面边距（GB/T 9704）
        self._set_page_margins()
        
        # 预热样式表（确保 Word 内置样式可用）
        self._warmup_styles()
        
        logger.info("[SemanticCompiler] 初始化完成（纯样式方案）")
    
    def _set_page_margins(self) -> None:
        """设置页面边距（GB/T 9704-2012）"""
        section = self.doc.sections[0]._sectPr
        margin = self.dfgp.page_margin
        
        pgMar = OxmlElement('w:pgMar')
        pgMar.set(qn('w:top'), str(margin.top_twips))
        pgMar.set(qn('w:bottom'), str(margin.bottom_twips))
        pgMar.set(qn('w:left'), str(margin.left_twips))
        pgMar.set(qn('w:right'), str(margin.right_twips))
        pgMar.set(qn('w:header'), str(int(0.7 * 567)))  # 7mm
        pgMar.set(qn('w:footer'), str(int(0.7 * 567)))
        pgMar.set(qn('w:gutter'), '0')
        
        section.append(pgMar)
        logger.debug("[SemanticCompiler] 页面边距已设置")
    
    def _warmup_styles(self) -> None:
        """
        预热样式表
        
        确保 Word 内置样式在文档中可用。
        python-docx 的 Document() 不会自动创建所有内置样式，
        我们需要通过访问来确保它们被创建。
        """
        try:
            for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
                _ = self.doc.styles[style_name]
            logger.debug("[SemanticCompiler] 样式表预热完成")
        except KeyError:
            # 样式不存在，创建基础定义
            logger.debug("[SemanticCompiler] 部分样式缺失，将使用 fallback")
    
    def build_from_ir(
        self, 
        ir_blocks: List[DocumentIRBlock], 
        output_path: str,
        toc_info: TocInfo = None
    ) -> None:
        """
        主循环：将 IR Block 序列重构为物理文档
        
        Args:
            ir_blocks: IR Block 序列
            output_path: 输出文件路径
            toc_info: 目录元数据（如有）
            
        Raises:
            CompilationError: 重构失败
        """
        if not ir_blocks:
            raise CompilationError("IR Block 序列为空")
        
        logger.info(f"[SemanticCompiler] 开始重构，共 {len(ir_blocks)} 个 Block")
        
        try:
            # 查找目录插入位置（如果有目录的话）
            toc_insert_idx = -1
            if toc_info and toc_info.has_toc:
                for i, block in enumerate(ir_blocks):
                    if block.label == 'SALUTATION':
                        toc_insert_idx = i + 1
                        break
            
            for i, block in enumerate(ir_blocks):
                self._process_block(block)
                
                # 在指定位置插入目录
                if i == toc_insert_idx and toc_info and toc_info.has_toc:
                    self._insert_toc(toc_info)
            
            self.doc.save(output_path)
            logger.info(f"[SemanticCompiler] 重构完成: {output_path}")
            
        except Exception as e:
            raise CompilationError(f"重构失败: {e}") from e
    
    def _insert_toc(self, toc_info: TocInfo) -> None:
        """插入目录（带分节符）"""
        toc_gen = TocGenerator()
        toc_gen.insert_toc(self.doc, toc_info, insert_idx=0)
        logger.info("[SemanticCompiler] 目录已插入")
    
    def _process_block(self, block: DocumentIRBlock) -> None:
        """
        处理单个 IR Block（纯样式方案）
        
        步骤：
        1. 获取该标签的格式参数
        2. 创建段落并写入纯文本（去除前导空格）
        3. 使用 python-docx API 设置样式
        4. 确保 outlineLevel 设置（TOC 生成必需）
        
        Args:
            block: IR Block
        """
        if not block.text:
            return
        
        # 去除前导空格（中文全角空格和ASCII空格）
        text = block.text.lstrip('\u3000 ')
        if not text:
            return
        
        # Step 1: 获取格式参数
        params = self.dfgp.get_style_params(block.label)
        
        # Step 2: 创建段落
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        
        # Step 3: 应用 Word 样式（优先使用 python-docx API）
        self._apply_word_style(paragraph, params)
        
        # Step 4: 设置中文字体（通过 run.font.name）
        self._apply_run_font(run, params)
        
        # Step 5: 设置段落格式（对齐、缩进、行距）
        self._apply_paragraph_format(paragraph, params)
        
        # Step 6: 确保 outlineLevel 设置（TOC 生成必需）
        self._ensure_outline_level(paragraph, params)
    
    def _apply_word_style(self, paragraph, params: StyleParams) -> None:
        """
        应用 Word 样式
        
        如果 params.word_style_name 存在且是 Word 内置样式，
        则使用 paragraph.style 设置样式。
        
        Args:
            paragraph: python-docx Paragraph 对象
            params: 样式参数
        """
        word_style = params.word_style_name
        if not word_style:
            return
        
        # 检查是否为 Word 内置样式
        if word_style not in BUILTIN_STYLE_NAMES:
            logger.debug(f"[Style] '{word_style}' 不是内置样式，跳过 style 设置")
            return
        
        # 检查样式是否在文档中可用
        try:
            style = self.doc.styles[word_style]
            # paragraph.style = style  # 禁用：避免蓝色主题字体
            logger.debug(f"[Style] 已应用 Word 样式: {word_style} (仅 outlineLevel，禁用style)")
        except KeyError:
            logger.debug(f"[Style] 样式 '{word_style}' 不存在，跳过")
    
    def _apply_run_font(self, run, params: StyleParams) -> None:
        """
        设置字体（纯样式方案）
        
        使用 python-docx run.font API 设置字体：
        - run.font.name: ASCII 字体
        - run.font.name_east_asia: 中文字体（东亚字体）
        
        Args:
            run: python-docx Run 对象
            params: 样式参数
        """
        font_family = params.font_family
        font_size_pt = params.font_size_pt
        
        # 直接操作 XML 设置字体（python-docx API 不够用）
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run._element.insert(0, rPr)
        
        # 查找或创建 rFonts
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        
        # 设置字体：ASCII/HAnsi=Times New Roman, eastAsia=中文字体
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), font_family)
        
        # 设置字号 (半磅单位)
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rPr.append(sz)
        sz.set(qn('w:val'), str(font_size_pt * 2))
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), str(font_size_pt * 2))
    
    def _apply_paragraph_format(self, paragraph, params: StyleParams) -> None:
        """
        设置段落格式（纯样式方案）
        
        使用 python-docx API 设置：
        - alignment: 对齐方式
        - first_line_indent: 首行缩进
        - right_indent: 右侧缩进
        - line_spacing: 行距
        
        Args:
            paragraph: python-docx Paragraph 对象
            params: 样式参数
        """
        # 1. 设置对齐方式
        if params.alignment:
            alignment_map = {
                'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
                'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
                'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            alignment = alignment_map.get(params.alignment.upper())
            if alignment:
                paragraph.alignment = alignment
                logger.debug(f"[Format] 对齐: {params.alignment}")
        
        # 2. 设置首行缩进
        if params.first_line_indent_twips is not None:
            try:
                paragraph.paragraph_format.first_line_indent = Twips(params.first_line_indent_twips)
                logger.debug(f"[Format] 首行缩进: {params.first_line_indent_twips} twips")
            except Exception as e:
                logger.debug(f"[Format] 设置首行缩进失败: {e}")
        
        # 3. 设置右侧缩进（"右空X字"）
        if params.right_indent_twips is not None:
            try:
                paragraph.paragraph_format.right_indent = Twips(params.right_indent_twips)
                logger.debug(f"[Format] 右侧缩进: {params.right_indent_twips} twips")
            except Exception as e:
                logger.debug(f"[Format] 设置右侧缩进失败: {e}")
        
        # 4. 设置行距
        if params.line_spacing_twips is not None:
            try:
                paragraph.paragraph_format.line_spacing = Twips(params.line_spacing_twips)
                paragraph.paragraph_format.line_spacing_rule = None  # 使用固定值
                logger.debug(f"[Format] 行距: {params.line_spacing_twips} twips (fixed)")
            except Exception as e:
                logger.debug(f"[Format] 设置行距失败: {e}")
        
        # 5. 设置段前段后间距
        if params.space_before_twips > 0:
            try:
                paragraph.paragraph_format.space_before = Twips(params.space_before_twips)
            except Exception as e:
                logger.debug(f"[Format] 设置段前间距失败: {e}")
        
        if params.space_after_twips > 0:
            try:
                paragraph.paragraph_format.space_after = Twips(params.space_after_twips)
            except Exception as e:
                logger.debug(f"[Format] 设置段后间距失败: {e}")
        
        # 6. 设置分页控制
        if params.keep_with_next:
            try:
                paragraph.paragraph_format.keep_with_next = True
            except Exception as e:
                logger.debug(f"[Format] 设置 keep_with_next 失败: {e}")
        
        if params.page_break_before:
            try:
                paragraph.paragraph_format.page_break_before = True
            except Exception as e:
                logger.debug(f"[Format] 设置 page_break_before 失败: {e}")
    
    def _ensure_outline_level(self, paragraph, params: StyleParams) -> None:
        """
        确保 outlineLevel 设置（TOC 生成必需）
        
        python-docx 的 paragraph.style 不总能正确设置 outlineLevel，
        因此我们通过底层 XML 确保 outlineLevel 被设置。
        
        Args:
            paragraph: python-docx Paragraph 对象
            params: 样式参数
        """
        if params.outline_level is None:
            return
        
        pPr = paragraph._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            paragraph._p.insert(0, pPr)
        
        # 检查是否已有 outlineLvl
        existing_outline = pPr.find(qn('w:outlineLvl'))
        if existing_outline is not None:
            pPr.remove(existing_outline)
        
        outlineLvl = OxmlElement('w:outlineLvl')
        outlineLvl.set(qn('w:val'), str(params.outline_level))
        pPr.append(outlineLvl)
        logger.debug(f"[TOC] outlineLevel: {params.outline_level}")


# === 便捷函数 ===

def compile_with_format(
    ir_blocks: List[DocumentIRBlock],
    output_path: str,
    dfgp_manager: Optional[DFGPManager] = None
) -> None:
    """
    一行代码：将 IR Block 序列编译为带格式的文档
    
    Args:
        ir_blocks: IR Block 序列
        output_path: 输出文件路径
        dfgp_manager: 可选 DFGP 管理器
    """
    compiler = SemanticCompiler(dfgp_manager)
    compiler.build_from_ir(ir_blocks, output_path)
