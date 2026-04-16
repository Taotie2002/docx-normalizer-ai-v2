"""
TocGenerator - 目录生成器
Phase 2 扩展模块

功能：
1. 在指定位置插入Word目录
2. 目录单独成页（前后分节符）
3. 设置页码重置（从1开始）
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from typing import Optional
import logging

from ..extractor.toc_detector import TocInfo

logger = logging.getLogger(__name__)


class TocGenerator:
    """
    目录生成器
    
    使用Word域代码生成目录：
    - TOC \\-o "1-3" \\-h \\-z \\-u  // 标题1-3，带超链接，隐藏页码，使用大纲级别
    """
    
    def __init__(self):
        self.toc_field_begin = 'TOC \\o "1-3" \\h \\z \\u'
    
    def insert_toc(self, doc: Document, toc_info: TocInfo, insert_idx: int) -> None:
        """
        在指定位置插入目录
        
        Args:
            doc: python-docx Document 对象
            toc_info: 目录元数据（TocInfo）
            insert_idx: 插入位置（段落索引）
        """
        if not toc_info.has_toc:
            logger.info("[TocGenerator] 文档无目录，不生成")
            return
        
        logger.info(f"[TocGenerator] 在 idx={insert_idx} 插入目录")
        
        # 获取插入位置的段落
        paragraphs = doc.paragraphs
        
        if insert_idx >= len(paragraphs):
            logger.warning(f"[TocGenerator] 插入位置 {insert_idx} 超出范围")
            insert_idx = len(paragraphs) - 1
        
        target_para = paragraphs[insert_idx]
        
        # 1. 在目录前插入分节符（下一页）
        self._insert_section_break_before(target_para)
        
        # 2. 创建目录标题段落
        toc_title = self._create_toc_title(doc, insert_idx)
        
        # 3. 在目录标题后插入TOC域
        self._insert_toc_field(doc, toc_title)
        
        # 4. 在目录后插入分节符（下一页）
        self._insert_section_break_after(doc, toc_title)
        
        # 5. 添加目录结束标记（空段落）
        self._insert_toc_end_marker(doc)
        
        # 6. 设置新节的页码（从1开始）
        self._reset_page_number(doc, toc_title)
    
    def _insert_section_break_before(self, para, break_type='nextPage') -> None:
        """
        在段落前插入分节符
        
        Args:
            para: 目标段落
            break_type: nextPage(下一页) / continuous(连续) / evenPage(偶数页) / oddPage(奇数页)
        """
        pPr = para._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para._p.insert(0, pPr)
        
        # 检查是否已有sectPr
        existing_sectPr = pPr.find(qn('w:sectPr'))
        if existing_sectPr is not None:
            pPr.remove(existing_sectPr)
        
        # 创建新的sectPr
        sectPr = OxmlElement('w:sectPr')
        
        # 页边距（继承当前节）
        pgMar = OxmlElement('w:pgMar')
        pgMar.set(qn('w:top'), '1800')      # 1.25cm
        pgMar.set(qn('w:right'), '1800')
        pgMar.set(qn('w:bottom'), '1800')
        pgMar.set(qn('w:left'), '1800')
        pgMar.set(qn('w:header'), '851')
        pgMar.set(qn('w:footer'), '851')
        pgMar.set(qn('w:gutter'), '0')
        sectPr.append(pgMar)
        
        # 分节类型
        sectType = OxmlElement('w:type')
        sectType.set(qn('w:val'), break_type)
        sectPr.append(sectType)
        
        # 页码格式
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:fmt'), 'decimal')
        pgNumType.set(qn('w:start'), '1')
        sectPr.append(pgNumType)
        
        pPr.append(sectPr)
        
        logger.debug(f"[TocGenerator] 插入分节符({break_type})于段落前")
    
    def _insert_section_break_after(self, doc: Document, after_para, break_type='nextPage') -> None:
        """
        在段落后插入分节符
        """
        # 创建空段落并设置sectPr
        new_para = OxmlElement('w:p')
        
        pPr = OxmlElement('w:pPr')
        sectPr = OxmlElement('w:sectPr')
        
        # 页边距
        pgMar = OxmlElement('w:pgMar')
        pgMar.set(qn('w:top'), '1800')
        pgMar.set(qn('w:right'), '1800')
        pgMar.set(qn('w:bottom'), '1800')
        pgMar.set(qn('w:left'), '1800')
        pgMar.set(qn('w:header'), '851')
        pgMar.set(qn('w:footer'), '851')
        pgMar.set(qn('w:gutter'), '0')
        sectPr.append(pgMar)
        
        # 分节类型
        sectType = OxmlElement('w:type')
        sectType.set(qn('w:val'), break_type)
        sectPr.append(sectType)
        
        # 页码重置为1
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:fmt'), 'decimal')
        pgNumType.set(qn('w:start'), '1')
        sectPr.append(pgNumType)
        
        pPr.append(sectPr)
        new_para.append(pPr)
        
        # 插入到after_para之后
        after_para._p.addnext(new_para)
        
        logger.debug(f"[TocGenerator] 插入分节符({break_type})于段落后")
    
    def _create_toc_title(self, doc: Document, idx: int) -> object:
        """
        创建目录标题段落
        """
        # 在idx位置插入目录标题
        toc_para = doc.add_paragraph()
        toc_para.alignment = 1  # CENTER
        
        # 设置字体为黑体
        run = toc_para.add_run('目    录')
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = doc.paragraphs[0].runs[0].font.size if doc.paragraphs[0].runs else None
        
        logger.debug("[TocGenerator] 创建目录标题: 目    录")
        return toc_para
    
    def _insert_toc_field(self, doc: Document, after_para: object) -> None:
        """
        插入TOC域代码
        
        Word TOC域: { TOC \\-o "1-3" \\-h \\-z \\-u }
        """
        # 创建包含域代码的段落
        toc_field_para = OxmlElement('w:p')
        
        # 域开始
        fldCharBegin = OxmlElement('w:fldChar')
        fldCharBegin.set(qn('w:fldCharType'), 'begin')
        run1 = OxmlElement('w:r')
        run1.append(fldCharBegin)
        toc_field_para.append(run1)
        
        # 域指令
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f' TOC \\o "1-3" \\h \\z \\u '
        run2 = OxmlElement('w:r')
        run2.append(instrText)
        toc_field_para.append(run2)
        
        # 域分隔
        fldCharSep = OxmlElement('w:fldChar')
        fldCharSep.set(qn('w:fldCharType'), 'separate')
        run3 = OxmlElement('w:r')
        run3.append(fldCharSep)
        toc_field_para.append(run3)
        
        # 占位符文本（Word更新后会替换）
        placeholder = OxmlElement('w:t')
        placeholder.text = '请右键单击此处或按 F9 更新目录'
        run4 = OxmlElement('w:r')
        run4.append(placeholder)
        toc_field_para.append(run4)
        
        # 域结束
        fldCharEnd = OxmlElement('w:fldChar')
        fldCharEnd.set(qn('w:fldCharType'), 'end')
        run5 = OxmlElement('w:r')
        run5.append(fldCharEnd)
        toc_field_para.append(run5)
        
        after_para._p.addnext(toc_field_para)
        
        logger.debug("[TocGenerator] 插入TOC域代码")
    
    def _insert_toc_end_marker(self, doc: Document) -> None:
        """
        插入目录结束标记（换行）
        """
        end_para = doc.add_paragraph()
        logger.debug("[TocGenerator] 插入目录结束标记")
    
    def _reset_page_number(self, doc: Document, after_toc_para: object) -> None:
        """
        确保正文节页码从1开始
        """
        # 在after_toc_para之后的段落设置sectPr的页码
        # 这在_insert_section_break_after中已经处理
        pass
